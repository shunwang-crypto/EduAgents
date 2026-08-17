"""Build the formatted Word version of 对话模块设计分享.md."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "对话模块设计分享.md"
OUTPUT = ROOT / "对话模块设计分享.docx"

TABLE_CAPTIONS = [
    "表 1  功能分层总览",
    "表 2  三个对话场景",
    "表 3  画像采集双路径分工",
    "表 4  个性化数据来源与作用",
    "表 5  关键设计决策与权衡",
]


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 10.5):
    run.font.name = latin
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "B7B7C9", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_inline(paragraph, text: str, *, italic_all: bool = False) -> None:
    """Render the small Markdown subset used by the source document."""
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run)
            run.italic = italic_all
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run)
            run.bold = True
            run.italic = italic_all
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, east_asia="等线", latin="Consolas", size=9.5)
            run.font.color.rgb = RGBColor(55, 55, 70)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run)
        run.italic = italic_all


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for name, size, color in (
        ("Heading 1", 16, "2F285E"),
        ("Heading 2", 13, "40367A"),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.paragraph_format.line_spacing = 1.4
        style.paragraph_format.space_after = Pt(2)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("对话模块设计分享 · 个性化实现")
    set_run_font(run, east_asia="微软雅黑", latin="Arial", size=9)
    run.font.color.rgb = RGBColor(105, 105, 120)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_run_font(run, size=9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = footer.add_run(" 页")
    set_run_font(tail, size=9)


def add_body_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.space_after = Pt(0)
    add_inline(paragraph, text)
    return paragraph


def add_caption(document: Document, text: str, *, italic: bool = False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    add_inline(paragraph, text, italic_all=italic)
    return paragraph


def add_table(document: Document, rows: list[list[str]], caption: str) -> None:
    add_caption(document, caption)
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index, cell in enumerate(row.cells):
            value = values[column_index] if column_index < len(values) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade_cell(cell, "E7E5F7" if row_index == 0 else ("F7F7FB" if row_index % 2 == 0 else "FFFFFF"))
            set_cell_border(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.2
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(9.5)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(47, 40, 94)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F1F1F5")
    p_pr.append(shd)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, east_asia="等线", latin="Consolas", size=9)


def add_image(document: Document, image_path: Path) -> None:
    max_width = 6.25
    max_height = 7.9
    image = DocxImage.from_file(str(image_path))
    width_px, height_px = image.px_width, image.px_height
    ratio = width_px / height_px
    width = max_width
    height = width / ratio
    if height > max_height:
        height = max_height
        width = height * ratio
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(4)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width), height=Inches(height))


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)

    table_index = 0
    heading1_count = 0
    in_code = False
    code_lines: list[str] = []
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()

        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not line:
            index += 1
            continue
        if line.startswith("# ") or line.startswith("> "):
            index += 1
            continue
        if line.startswith("## "):
            if heading1_count:
                document.add_page_break()
            paragraph = document.add_paragraph(style="Heading 1")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline(paragraph, line[3:].strip())
            heading1_count += 1
            index += 1
            continue
        if line.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline(paragraph, line[4:].strip())
            index += 1
            continue
        if line.startswith("!["):
            match = re.match(r"!\[[^]]*\]\(([^)]+)\)", line)
            if match:
                add_image(document, (ROOT / match.group(1)).resolve())
            index += 1
            continue
        if line.startswith("*") and line.endswith("*") and line[1:-1].startswith("图"):
            add_caption(document, line[1:-1], italic=True)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[index + 1].strip()):
            table_lines = [line]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [[cell.strip() for cell in table_line.strip("|").split("|")] for table_line in table_lines]
            caption = TABLE_CAPTIONS[table_index] if table_index < len(TABLE_CAPTIONS) else f"表 {table_index + 1}"
            add_table(document, rows, caption)
            table_index += 1
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline(paragraph, line[2:].strip())
            index += 1
            continue
        ordered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if ordered:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline(paragraph, ordered.group(2))
            index += 1
            continue

        add_body_paragraph(document, line)
        index += 1

    if in_code and code_lines:
        add_code_block(document, code_lines)

    core = document.core_properties
    core.title = "对话模块设计分享：个性化实现"
    core.subject = "EduAgents 对话模块、个性化上下文与安全记忆机制"
    core.author = "EduAgents 项目组"
    core.keywords = "EduAgents, 个性化, 学习者画像, 对话, 记忆, RAG"

    temporary = OUTPUT.with_suffix(".tmp.docx")
    document.save(temporary)
    temporary.replace(OUTPUT)


if __name__ == "__main__":
    build()
